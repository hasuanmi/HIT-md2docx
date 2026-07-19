#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
cover_inject.py — HIT 硕士论文"封面与前置部分版式注入"（标准后处理步骤，不改动引擎）

职责（全部依赖外部官方模板 封面.docx，引擎本身不生成这种精确版式）：
  1. 整页复制 封面.docx 的封面区（含 3 页：中文封面 / 中文扉页表格 / 英文封面表格 +
     分页符 + 分节 sectPr）到论文开头，替换其中的中/英文题目为你论文的题目，
     白色批注箭头原样保留。
  2. 在中/英文目录的章节条目之前，插入 `摘  要 → I`、`Abstract → II`（罗马数字页码）。
  3. 对齐前置部分（封面 + 摘要页）行距到官方《书写范例》：摘要标题段距、摘要正文行距。

题目来源（按优先级）：
  a) 命令行 --cn-title / --en-title（agent 在 2.5 预翻译后传入，最高优先）
  b) front_matter（论文自己的"论文题目："/"英文题目："；模板不应写死真实题目）
  c) 从输入 docx 现有封面区提取（自包含兜底）
  d) 英文题目在 a/b/c 均缺失、且配置了 HITMD2DOCX_LLM_API_KEY 时，
     由 thesis_md2docx.translation 调用 LLM 翻译中文题目生成（与英文目录
     同一套机制，且受同一个 HITMD2DOCX_NO_LLM 开关控制），绝不回退写死样例。

健壮性：
  - 找不到 封面.docx 模板 → 跳过封面复制，仅做目录/行距，不报错（流程可无模板运行）。
  - 幂等：每次都基于模板重新注入封面区（先删旧封面区）；目录条目/行距检测已存在则跳过。

用法：
  python cover_inject.py <thesis.docx> [--cover-template input/封面.docx]
                              [--front-matter input/front_matter_hit.md]
                              [--out output/<论文>.docx] [--no-spacing]
"""
from __future__ import annotations
import argparse
import copy
import json
import os
import re
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

import math

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
Q = lambda t: f"{{{W}}}{t}"

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ENGINE_ROOT = os.path.dirname(os.path.dirname(SCRIPT_DIR))   # HIT-md2docx/ (scripts -> skill -> HIT-md2docx)
DEFAULT_COVER = os.path.join(ENGINE_ROOT, "input", "封面.docx")
DEFAULT_FRONT = os.path.join(ENGINE_ROOT, "input", "front_matter_hit.md")


# ---------------------------------------------------------------------------
# 题目提取
# ---------------------------------------------------------------------------
def _read_front_titles(front_path: str):
    cn = en = None
    if front_path and os.path.exists(front_path):
        with open(front_path, encoding="utf-8") as f:
            for line in f:
                s = line.strip()
                if s.startswith("论文题目：") or s.startswith("论文题目:"):
                    cn = s.split("：", 1)[-1].split(":", 1)[-1].strip()
                elif s.startswith("英文题目：") or s.startswith("英文题目:"):
                    en = s.split("：", 1)[-1].split(":", 1)[-1].strip()
    return cn, en


# 占位符特征：模板里写死的「（请替换…）/【请填写】/To Be Filled」都算未提供
_PLACEHOLDER_HINTS = ("请替换", "待填写", "To Be Filled", "Please Fill",
                      "示例", "样例", "xxx", "XXX", "（请", "【请")
def _is_placeholder(val: str | None) -> bool:
    if not val:
        return True
    v = val.strip()
    if not v:
        return True
    if v.startswith("（") or v.startswith("("):
        return True
    if v.startswith("【") or v.startswith("["):
        return True
    return any(h in v for h in _PLACEHOLDER_HINTS)


# 正文首个 H1 题目：跳过「第X章」、以数字开头的章号、摘要/ABSTRACT
_MD_H1_TITLE = re.compile(r"^#\s+(?!第[一二三四五六七八九十零\d]+章)(?!\d)(.+)$")
_MD_TITLE_SKIP = {"摘要", "ABSTRACT", "Abstract"}


def _read_md_title(md_path: str):
    """从（预处理后的）论文 md 取题名，作封面题名兜底。

    优先级：正文里的 `论文题目：`/`中文题目：` 行 → 第一个 H1 标题
    （用户最常见写法就是把题名写成首行 `# 《…》…`）。取不到返回 (None, None)。
    """
    if not md_path or not os.path.exists(md_path):
        return None, None
    try:
        with open(md_path, encoding="utf-8") as f:
            lines = f.readlines()
    except Exception:
        return None, None
    cn = en = None
    for s in (ln.strip() for ln in lines):
        if cn is None and (s.startswith("论文题目：") or s.startswith("论文题目:")):
            cn = s.split("：", 1)[-1].split(":", 1)[-1].strip()
        elif cn is None and (s.startswith("中文题目：") or s.startswith("中文题目:")):
            cn = s.split("：", 1)[-1].split(":", 1)[-1].strip()
        elif en is None and (s.startswith("英文题目：") or s.startswith("英文题目:")):
            en = s.split("：", 1)[-1].split(":", 1)[-1].strip()
    if cn is None:
        for s in (ln.strip() for ln in lines):
            m = _MD_H1_TITLE.match(s)
            if m:
                cand = m.group(1).strip()
                if cand in _MD_TITLE_SKIP:
                    continue
                cn = cand
                break
    return cn, en


def _read_json_title(markdown_dir, cn):
    """从 <md同目录>/heading_translations.json 查封面英文题名（agent 预翻译）。

    与英文目录同一份字典：agent 在 Step 2.5 已将「论文总标题 → 英文」
    翻好写回该文件（见 generate.py --dump-headings 对总标题的导出）。
    这里用归一化后的中文总题 key 精确命中，使封面英文题也能零 key 自动翻译。
    找不到文件 / key 未命中 / 值为空时返回 None，交由 docx 提取或 LLM 兜底。
    """
    if not cn or not markdown_dir:
        return None
    jp = Path(markdown_dir) / "heading_translations.json"
    if not jp.exists():
        return None
    try:
        data = json.loads(jp.read_text(encoding="utf-8"))
    except Exception:
        return None
    try:
        from thesis_md2docx.toc import normalize_heading_key
    except Exception:
        def normalize_heading_key(x):
            return (x or "").strip()
    key = normalize_heading_key(cn)
    for k, v in data.items():
        if normalize_heading_key(k) == key and v:
            return v
    return None


def _extract_titles_from_doc(doc: Document):
    """从 docx 现有封面区（开头到第一个摘要）提取中/英文题目，作兜底。"""
    cn = en = en_lines = None
    cn_cands, en_cands = [], []
    FIXED = ("硕士学位论文", "学术学位", "哈尔滨工业大学", "国内图书分类号",
             "国际图书分类号", "学校代码", "密级", "Classified", "U.D.C",
             "Dissertation", "Candidate", "Supervisor", "Academic Degree",
             "Speciality", "Affiliation", "Date of Defence",
             "Degree-Conferring", "授予学位", "所在单位", "答辩日期", "申请学位", "导师")
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if "摘" in t and "要" in t and "Abstract" not in t:
            break  # 到达摘要，封面区结束
        if any(k in t for k in FIXED):
            continue
        if re.search(r"[一-鿿]", t) and len(t) >= 8 and not t.startswith("（"):
            cn_cands.append(t)
        if re.match(r"^[A-Z][a-z]", t) and ("Resilience" in t or "Research" in t
                                            or "Convergence" in t or "Impact" in t):
            en_cands.append(t)
    if cn_cands:
        cn = cn_cands[0]
    if en_cands:
        en = " ".join(en_cands).replace("↑", "").strip()
    return cn, en


def _llm_translate_title(cn: str, markdown_dir=None):
    """与英文目录 translation.py 同一套 LLM 兜底：配置了 API key 时翻译中文题目。

    返回英文题字符串或 None（未配置 key / 翻译失败 / 被 HITMD2DOCX_NO_LLM 禁用）。
    绝不返回写死的样例标题。
    """
    if not cn or os.environ.get("HITMD2DOCX_NO_LLM"):
        return None
    try:
        if ENGINE_ROOT not in sys.path:
            sys.path.insert(0, ENGINE_ROOT)
        from thesis_md2docx import translation
    except Exception:
        return None
    try:
        cfg = translation.load_llm_config(
            Path(markdown_dir) if markdown_dir else None
        )
    except Exception:
        return None
    if not cfg:
        return None
    try:
        tr = translation.translate_headings([cn], config=cfg)
        return tr.get(cn)
    except Exception:
        return None


def extract_titles(front_path, doc, cli_cn=None, cli_en=None, markdown_dir=None, md_path=None):
    fcn, fen = _read_front_titles(front_path)
    mcn, men = _read_md_title(md_path)
    # 占位符视为「未提供」，避免覆盖正文真实题名（默认无 --front-matter
    # 时 front_path 是写死占位符的模板，必须让位给正文 md 的首行 H1）
    if _is_placeholder(fcn):
        fcn = None
    if _is_placeholder(fen):
        fen = None
    if _is_placeholder(mcn):
        mcn = None
    if _is_placeholder(men):
        men = None
    # 优先级：CLI > front_matter > 正文 md（首行 H1 / 论文题目：行）> docx 提取
    cn = cli_cn or fcn or mcn
    en = cli_en or fen or men
    # 新增：agent 在 Step 2.5 已把「论文总标题 → 英文」也翻好写回同一份
    # heading_translations.json（见 generate.py --dump-headings 对总标题的导出）。
    # 用归一化中文总题精确命中，实现封面英文题零 key 自动翻译，与英文目录同源。
    if not en and cn:
        en = _read_json_title(markdown_dir, cn)
    if not cn or not en:
        dcn, den = _extract_titles_from_doc(doc)
        cn = cn or dcn
        en = en or den
    # 与英文目录一致：a/b/c 都缺英文题时，配了 LLM key 则调用 LLM 生成，
    # 绝不回退到写死的样例英文题（模板里的「英文题目：」应为空）。
    if not en and cn:
        en = _llm_translate_title(cn, markdown_dir)
    return cn, en


# ---------------------------------------------------------------------------
# 题目替换
# ---------------------------------------------------------------------------
def _replace_title_text(para, new_text):
    """替换段落文本，并清空其余 run，避免模板残留的格式说明/乱码符号被保留。

    只改文本，绝不改字号/字体/间距（题目字号须严格保持模板规范：2号字）。
    """
    runs = para.runs
    if not runs:
        return False
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""
    return True


def _split_en(en, maxc=56):
    """把英文题目按模板两行结构拆分。

    如果整题能在 56 字符以内放下，则只放第一行，第二行留空；
    否则优先在第一行结束时断开，剩余全部放入第二行。
    """
    if len(en) <= maxc:
        return en, ""
    words = en.split()
    if len(words) <= 1:
        return en, ""
    line1 = words[0]
    j = 1
    while j < len(words) and len(line1) + 1 + len(words[j]) <= maxc:
        line1 += " " + words[j]
        j += 1
    return line1, " ".join(words[j:])


# 封面固定标签（非题目），用于从模板中识别并排除，避免误把标签当题目替换
_COVER_FIXED_LABELS = (
    "哈尔滨工业大学", "本科学位论文", "硕士学位论文", "博士学位论文",
    "国内图书分类号", "国际图书分类号", "学校代码", "密级",
    "Classified Index", "U.D.C", "Dissertation for the Master Degree",
    "Candidate", "Supervisor", "Academic Degree Applied for", "Speciality",
    "Affiliation", "Date of Defence", "Degree-Conferring", "授予学位单位",
    "所在单位", "申请学位", "导师", "答辩日期", "Harbin Institute of Technology",
    "（Times New Roman", "Times New Roman", "2号", "小2号", "加粗", "居中",
    "年 月", "年月", "答辩日期",
)


def _is_cover_label(t: str) -> bool:
    return any(k in t for k in _COVER_FIXED_LABELS)


def _para_is_bold_title_style(p) -> bool:
    """判断段落是否符合封面题目（中文/英文）的格式特征：至少有一个非白色 run 加粗，
    且字号 >= 18pt（小二号）= 342000 EMU。"""
    for r in p.runs:
        if not r.text:
            continue
        # 跳过白色批注 run
        rPr = r._r.find(Q("rPr"))
        if rPr is not None:
            c = rPr.find(Q("color"))
            if c is not None and c.get(Q("val"), "").upper() == "FFFFFF":
                continue
        if r.bold:
            sz = r.font.size
            if sz is not None and sz >= 342000:  # 18pt
                return True
    return False


def _cjk_len(t: str) -> int:
    return sum(1 for c in t if "一" <= c <= "鿿")


def _en_len(t: str) -> int:
    return sum(1 for c in t if c.isascii() and c.isalpha())


# ---------------------------------------------------------------------------
# 封面第一页排版守恒：题目变长后压缩空行，保证"哈尔滨工业大学 / 年 月"留在第一页底部
#
# 思路（不改字号，严格保持模板 2 号题字规范）：
#   模板用一串空行把校名/年月块顶到第一页底部。题目越长、换行越多，就把校名块
#   往下挤到第二页。因此按"题目相对模板原题多出的视觉行数"，等比例删掉题目区与
#   校名块之间的空行即可把校名块拉回第一页底部。
#   - 题目行(2号=22pt) 行高 ≈ 27.5pt；空行(正文五号=10.5pt) 行高 ≈ 12.6pt。
#   - 每多 1 行题目 ≈ 需删 27.5/12.6 ≈ 2.2 个空行来补偿。
# ---------------------------------------------------------------------------
_TITLE_LINE_PT = 27.5      # 2号题目单行高度(近似)
_EMPTY_LINE_PT = 12.6      # 五号空行单行高度(近似)
_TITLE_FONT_PT = 22.0      # 2号
_EN_CHAR_W_RATIO = 0.5     # 英文字符平均宽 ≈ 0.5×字号


def _cover_usable_width_pt(cover: Document) -> float:
    """封面页面可用文字宽度(pt) = 页宽 - 左右页边距。"""
    body = cover.element.body
    for sect in body.iter(Q("sectPr")):
        pgSz = sect.find(Q("pgSz"))
        pgMar = sect.find(Q("pgMar"))
        if pgSz is not None and pgMar is not None:
            try:
                w = int(pgSz.get(Q("w")))
                l = int(pgMar.get(Q("left")))
                r = int(pgMar.get(Q("right")))
                return (w - l - r) / 20.0  # twips -> pt
            except (TypeError, ValueError):
                pass
    return 425.0  # A4 默认(11906-1701-1701)/20


def _text_width_units(t: str) -> float:
    """估算文本"全角宽度单位"：中日韩/全角标点=1.0，其余(含 ASCII)=0.5。"""
    u = 0.0
    fullwidth_punct = "《》【】「」『』（）—…·、，。；：！？"
    for c in t:
        if ("\u4e00" <= c <= "\u9fff") or ("\u3000" <= c <= "\u303f") \
                or ("\uff00" <= c <= "\uffef") or c in fullwidth_punct:
            u += 1.0
        else:
            u += 0.5
    return u


def _cjk_title_lines(text: str, usable_pt: float) -> int:
    """中文题目在封面题目区的视觉行数。"""
    if not text:
        return 0
    units = _text_width_units(text)
    per_line = max(1.0, usable_pt / _TITLE_FONT_PT)
    return max(1, math.ceil(units / per_line))


def _en_title_lines(text: str, usable_pt: float) -> int:
    """英文题目在封面题目区的视觉行数。"""
    n = len((text or "").strip())
    if not n:
        return 0
    per_line = max(1.0, usable_pt / (_TITLE_FONT_PT * _EN_CHAR_W_RATIO))
    return max(1, math.ceil(n / per_line))


def shrink_cover_remaining_gaps(cover: Document, target_pt: float = 6.0) -> int:
    """将「哈尔滨工业大学」上方剩余空行的行距缩小到 target_pt（兜底机制）。

    当题目特别长、仅靠删除空行仍无法把年月拉回首页时，进一步压缩
    留存空行的高度。每行约省 (12.6 - target_pt) pt。
    """
    paras = cover.paragraphs
    anchor = None
    for i, p in enumerate(paras):
        if "哈尔滨工业大学" in p.text:
            anchor = i
            break
    if anchor is None:
        return 0
    n = 0
    j = anchor - 1
    while j >= 0 and paras[j].text.strip() == "":
        _set_para_spacing(paras[j], line=int(target_pt * 20), rule="auto")  # pt → twips (1pt=20twips)
        n += 1
        j -= 1
    return n


def compress_cover_gap(cover: Document, remove_n: int, keep_min: int = 2) -> int:
    """删除"哈尔滨工业大学"上方紧邻的若干空行，把校名/年月块拉回第一页底部。

    保护块 = 「哈尔滨工业大学」段 + 其后紧跟的「年 月」段（两者必须在同一页）。
    只删除紧邻校名块上方的连续空段落（题目区与校名块之间的排版留白），
    保留至少 keep_min 个空行以维持底部间距；从最靠近校名块的一侧开始删。
    """
    if remove_n <= 0:
        return 0
    paras = cover.paragraphs
    anchor = None
    year_month_idx = None
    for i, p in enumerate(paras):
        if "哈尔滨工业大学" in p.text:
            anchor = i
            # 向下找紧跟的「年 月」段（模板中两者相邻）
            j = i + 1
            while j < len(paras) and paras[j].text.strip() == "":
                j += 1
            if j < len(paras) and ("年" in paras[j].text and "月" in paras[j].text):
                year_month_idx = j
            break
    if anchor is None:
        return 0

    # 空行列表：校名段上方连续空段
    empties = []
    j = anchor - 1
    while j >= 0 and paras[j].text.strip() == "":
        empties.append(paras[j])
        j -= 1

    # 如果年月段与校名段之间还有空行，也纳入可压缩范围
    if year_month_idx is not None and year_month_idx > anchor + 1:
        for k in range(anchor + 1, year_month_idx):
            if paras[k].text.strip() == "":
                empties.append(paras[k])

    removable = max(0, len(empties) - keep_min)
    to_remove = min(remove_n, removable)
    removed = 0
    for k in range(to_remove):
        el = empties[k]._p
        el.getparent().remove(el)
        removed += 1
    return removed


def replace_cover_titles(cover: Document, cn, en):
    """把模板封面的中/英文题目替换为论文题目。

    定位策略基于「题目 = 封面中长度达标的非固定标签文本段落」：
      - 中文题目：替换所有 CJK 长度 >= 8、非固定标签、非格式说明的段落；
      - 英文题目：替换所有纯英文长度 >= 20、非固定标签、非格式说明的段落，
        并把其后续连续英文续行清空或填入第二行。
    模板底部的格式说明（如「年 月、Times New Roman...」）已被 _COVER_FIXED_LABELS
    排除，避免被误替换。
    """
    if not cn and not en:
        return 0
    paras = cover.paragraphs
    n = 0

    # —— 替换前先记录模板原题目，用于估算题目变长后多出的视觉行数 ——
    usable_pt = _cover_usable_width_pt(cover)
    orig_cn = None
    for p in paras:
        t = p.text.strip()
        if not t or _is_cover_label(t) or t.lstrip().startswith("（"):
            continue
        if _cjk_len(t) >= 8:
            orig_cn = t
            break
    orig_en_parts = []
    for i0 in range(len(paras)):
        t = paras[i0].text.strip()
        if t and not _is_cover_label(t) and _en_len(t) >= 20 and _cjk_len(t) == 0:
            orig_en_parts.append(t)
            k = i0 + 1
            while k < len(paras):
                tt = paras[k].text.strip()
                if not tt:
                    k += 1
                    continue
                if _is_cover_label(tt):
                    break
                if _en_len(tt) > 0 and _cjk_len(tt) == 0:
                    orig_en_parts.append(tt)
                    k += 1
                    continue
                break
            break
    orig_en = " ".join(orig_en_parts)

    if cn:
        for p in paras:
            t = p.text.strip()
            if not t or _is_cover_label(t) or t.lstrip().startswith("（"):
                continue
            if _cjk_len(t) >= 8:
                if _replace_title_text(p, cn):
                    n += 1

    if en:
        en1, en2 = _split_en(en)
        i = 0
        while i < len(paras):
            p = paras[i]
            t = p.text.strip()
            if not t or _is_cover_label(t):
                i += 1
                continue
            if _en_len(t) >= 20 and _cjk_len(t) == 0:
                # 命中英文标题第一行
                if _replace_title_text(p, en1):
                    n += 1
                # 处理续行：下一个连续英文段落作为第二行
                j = i + 1
                second_filled = False
                while j < len(paras):
                    tt = paras[j].text.strip()
                    if not tt:
                        j += 1
                        continue
                    if _is_cover_label(tt):
                        break
                    if _en_len(tt) > 0 and _cjk_len(tt) == 0:
                        if en2 and not second_filled:
                            if _replace_title_text(paras[j], en2):
                                n += 1
                            second_filled = True
                        else:
                            for r in paras[j].runs:
                                r.text = ""
                        j += 1
                        continue
                    # 遇到非英文格式说明行，停止清空
                    break
                i = j
            else:
                i += 1

    # —— 题目变长 → 适量压缩校名块上方空行，把"哈尔滨工业大学 / 年 月"拉回第一页底部 ——
    # 注意：原模板（17字+79字符英文）在A4封面下恰好让校名落在第1页底部（15个空行）。
    # 用户题目（22字+109字符）算法估多 1 行 → 实际渲染多 0~1 行。
    # 不删/微删即可，不要过度压缩（过度会把校名挤到第2页）。
    # 布局锚定模式：封面模板中「哈尔滨工业大学」段已带 w:framePr（固定钉在 P1 底部），
    # 校名/年月与标题行数彻底解耦，无需也不会压缩空行，直接跳过。
    anchored = False
    for p in paras:
        if "哈尔滨工业大学" in p.text:
            pPr = p._p.find(Q("pPr"))
            if pPr is not None and pPr.find(Q("framePr")) is not None:
                anchored = True
            break
    if anchored:
        print("   布局锚定模式：校名/年月已固定钉在 P1 底部，跳过空行压缩")
        return n
    base_lines = _cjk_title_lines(orig_cn or "", usable_pt) + _en_title_lines(orig_en, usable_pt)
    new_lines = _cjk_title_lines(cn or orig_cn or "", usable_pt) + _en_title_lines(en or orig_en, usable_pt)
    extra_lines = new_lines - base_lines
    if extra_lines > 0:
        # 只删除实际多出的视觉行对应的高度（用更保守的换算：1行标题 ≈ 1.5 行空行）
        # 不加额外安全余量，否则会把校名挤到第2页
        remove_n = max(1, round(extra_lines * 1.5))
        removed = compress_cover_gap(cover, remove_n)
        print(f"   题目多出 {extra_lines} 行 → 压缩封面空行 {removed} 个（保校名+年月于首页底部）")

    return n


# ---------------------------------------------------------------------------
# 封面区元素构建 & 注入
# ---------------------------------------------------------------------------
def _strip_foreign_attributes(elems):
    """删除 cover elems 中属于"非核心"命名空间的属性（如 w14:paraId / w14:docId）。

    这些是 Word 给段落/运行自动添加的内部 ID，与版式渲染无关；删除后 Word
    在打开文档时会自动重建。关键作用是：避免把带 w14:* 属性的封面元素 deepcopy
    注入到「根 <w:document> 未声明 w14 命名空间」的引擎产物后，Word 因遇到未声明
    前缀而报"文件可能已经损坏"。

    保留的命名空间：w(正文) / r(关系，如 r:id) / xml(xml:space 等)。
    """
    KEEP = {
        W,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "http://www.w3.org/XML/1998/namespace",
    }
    def _walk(e):
        if not isinstance(e.tag, str):
            return
        for k in list(e.attrib.keys()):
            if not isinstance(k, str):
                continue
            ns = etree.QName(k).namespace
            if ns not in KEEP:
                del e.attrib[k]
        for c in e:
            _walk(c)
    for e in elems:
        _walk(e)


def build_cover_elems(cover: Document):
    """返回封面模板 body 下所有子元素的深拷贝，并保留其原始分节结构。

    封面模板 ``input/封面.docx`` 的 body 末尾会带一个 ``w:sectPr``（body sectPr）。
    在目标文档中，这个 sectPr 若作为 body 子元素直接插入到封面与正文之间，会
    破坏分节逻辑，导致英文封面与摘要之间的空白页丢失或分节错位。

    这里把它从 body 子元素中取出，作为 inline sectPr 挂到封面区最后一个段落
    （即空白页段落）的 pPr 上：
      - 封面区最后一个 inline sectPr（模板中已有的英文封面节结束）会先产生新页，
        进入空白页；
      - 移到这个末段的 body sectPr 再产生新页，使摘要从新页开始；
      - 这样英文封面与摘要之间自然夹着一页空白页。
    """
    cover_body = cover.element.body
    children = list(cover_body)
    elems = [copy.deepcopy(c) for c in children]

    # 剥掉外部命名空间属性（w14:paraId 等），避免注入后 Word 报损坏
    _strip_foreign_attributes(elems)

    # 把模板末尾的 body sectPr 取下来，准备挂到封面区最后一个段落
    last_sectPr = None
    if elems and etree.QName(elems[-1]).localname == "sectPr":
        last_sectPr = elems.pop()

    # 确保封面区最后一个元素是段落（模板正常是末尾的空白页段落）
    if not elems or etree.QName(elems[-1]).localname != "p":
        blank = OxmlElement("w:p")
        elems.append(blank)
    last_p = elems[-1]

    if last_sectPr is not None:
        pPr = last_p.find("{%s}pPr" % W)
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            last_p.insert(0, pPr)
        # 避免和已有 sectPr 冲突：先移除已有的 inline sectPr
        old = pPr.find("{%s}sectPr" % W)
        if old is not None:
            pPr.remove(old)
        pPr.append(last_sectPr)

    return elems


def inject_cover(target: Document, cover_elems):
    body = target.element.body
    cut = None
    for child in body:
        if etree.QName(child).localname == "p":
            txt = "".join(child.itertext()).strip()
            if "摘" in txt and "要" in txt and "Abstract" not in txt and "ABSTRACT" not in txt.upper():
                cut = child
                break
    if cut is None:
        print("  [跳过封面替换] 未找到摘要标题，不删除封面区")
        return False

    to_remove = []
    for child in body:
        if child is cut:
            break
        to_remove.append(child)
    for c in to_remove:
        body.remove(c)
    for el in reversed(cover_elems):
        body.insert(0, el)

    # —— 确保中文扉页与英文封面之间存在空白页 ——
    # 锚定模板里靠 5 个空段自然撑出空白页；但当封面接在长文档前面时，
    # Word 的全局 Repaginate 可能因后续分节干扰而合并页面。
    # 这里做确定性保护：若检测到「中文扉页表格」和「Classified Index」之间没有
    # nextPage 分节符 / 硬分页，则插入一个 nextPage 分节符强制空白页。
    _ensure_blank_page_between_cn_en_covers(body)

    return True


def _ensure_blank_page_between_cn_en_covers(body):
    """在中文扉页(含信息表)与英文封面(Classified Index)之间确保有一张空白页。

    机制（确定性，不依赖溢出撑页）：
      中文扉页结束 → 插入一个"硬分页空段"(<w:br w:type="page"/>) → 英文封面首段
      自带 <w:br w:type="page"/>。两道硬分页之间恰好夹着一页空白页。

    插入的是"分页空段"而非 nextPage 分节符——前者不创建新节、不触碰
    页眉/页脚继承，对 Word 兼容性更稳妥；后者会新建一节，在部分严格
    解析环境下可能触发误报。
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    def q(t): return f"{{{W}}}{t}"

    children = list(body)
    cn_end = -1       # 中文扉页表格末尾索引
    en_start = -1     # 英文封面起始索引(Classified Index)

    for i, el in enumerate(children[:120]):  # 封面区通常在前 120 个元素内
        ln = etree.QName(el).localname
        if ln == "p":
            t = "".join(el.itertext()).strip()
            if ("授予学位单位" in t) or ("哈尔滨工业大学" in t and "分类" not in t and "Date" not in t):
                cn_end = i
            if "Classified Index" in t:
                en_start = i
                break
        elif ln == "tbl":
            txt = "".join(el.itertext())
            if "导师" in txt or "学科" in txt:
                cn_end = i

    if cn_end < 0 or en_start < 0 or cn_end >= en_start:
        return  # 未找到预期结构，不改动

    # 是否已在中英文封面"之间"(不含英文封面自身)存在硬分页 → 已保证空白页，跳过
    def _has_page_break(el):
        if etree.QName(el).localname != "p":
            return False
        pPr = el.find(q("pPr"))
        if pPr is not None:
            if pPr.find(q("sectPr")) is not None:
                return True
            if pPr.find(q("pageBreakBefore")) is not None:
                return True
        for br in el.iter(q("br")):
            if br.get(q("type")) == "page":
                return True
        return False

    for i in range(cn_end + 1, en_start):
        if _has_page_break(children[i]):
            return  # 已有分页保护，空白页已保证

    # 英文封面自身是否带硬分页（自带 br=page 或 pageBreakBefore）
    en_has_pb = _has_page_break(children[en_start]) if en_start < len(children) else False

    # 在中文扉页后插入硬分页空段
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(q("type"), "page")
    r.append(br)
    p.append(r)
    body.insert(cn_end + 1, p)

    # 若英文封面自身没有硬分页，则给它加 pageBreakBefore，确保两段分页 = 一张空白页
    if not en_has_pb and en_start < len(children):
        en_el = children[en_start]
        if etree.QName(en_el).localname == "p":
            pPr = en_el.find(q("pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                en_el.insert(0, pPr)
            pbb = OxmlElement("w:pageBreakBefore")
            pPr.append(pbb)

    print(f"   已在中文扉页与英文封面之间插入硬分页空段，保证一张空白页")


# ---------------------------------------------------------------------------
# 目录摘要条目
# ---------------------------------------------------------------------------
def _find_toc_title(paras, labels):
    for i, p in enumerate(paras):
        t = p.text.strip().replace(" ", "")
        if t in labels:
            return i
    return None


def _first_chapter(paras, start):
    for i in range(start, len(paras)):
        t = paras[i].text.strip()
        if re.match(r"^第.+章", t) or re.match(r"^\d+\s", t) or re.match(r"^[一二三四五六七八九十]+、", t):
            return i
    return None


def _first_chapter_en(paras, start):
    for i in range(start, len(paras)):
        if paras[i].text.strip().startswith("Chapter"):
            return i
    return None


def _has_entry(paras, start, end, text):
    for i in range(start, end):
        if paras[i].text.strip().startswith(text):
            return True
    return False


def _make_entry_para(before_para, text, page):
    p = OxmlElement("w:p")
    pPr = before_para._p.find(Q("pPr"))
    if pPr is not None:
        p.append(copy.deepcopy(pPr))
    r1 = OxmlElement("w:r")
    t1 = OxmlElement("w:t")
    t1.set(Q("space"), "preserve")
    t1.text = text
    r1.append(t1)
    rt = OxmlElement("w:r")
    rt.append(OxmlElement("w:tab"))
    r2 = OxmlElement("w:r")
    t2 = OxmlElement("w:t")
    t2.set(Q("space"), "preserve")
    t2.text = page
    r2.append(t2)
    p.append(r1)
    p.append(rt)
    p.append(r2)
    return p


def _insert_toc_entries(target, before_para, entries):
    body = target.element.body
    bp = before_para._p
    idx = list(body).index(bp)
    for text, page in reversed(entries):
        p = _make_entry_para(before_para, text, page)
        body.insert(idx, p)


def _is_toc_para(p) -> bool:
    """判断段落是否为目录条目（带 TOC1/2/3 样式或分页字段）。"""
    pPr = p._p.find(Q("pPr"))
    if pPr is None:
        return False
    ps = pPr.find(Q("pStyle"))
    if ps is not None:
        style_id = ps.get(Q("val")) or ""
        if style_id.upper().startswith("TOC"):
            return True
    # 含 TOC 分页字段也视为目录条目
    for r in p._p.findall(Q("r")):
        for fc in r.iter(Q("fldChar")):
            return True
    return False


def ensure_toc_abstract(target: Document):
    paras = target.paragraphs
    n = 0
    ci = _find_toc_title(paras, ("目  录", "目录", "目 录"))
    if ci is not None:
        fc = _first_chapter(paras, ci + 1)
        if fc is not None and not _has_entry(paras, ci, fc, "摘  要"):
            _insert_toc_entries(target, paras[fc], [("摘  要", "I"), ("Abstract", "II")])
            n += 2
    # 英文目录（Contents）由引擎直接生成 Abstract(In Chinese) / Abstract(In English)，
    # 不再插入 摘 要/Abstract 中文标签，避免重复。
    return n


# ---------------------------------------------------------------------------
# 前置部分行距对齐（官方书写范例）
# ---------------------------------------------------------------------------
def _set_para_spacing(p, before=None, after=None, line=None, rule=None):
    pPr = p._p.get_or_add_pPr()
    sp = pPr.find(Q("spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    if before is not None:
        sp.set(Q("before"), str(int(before)))
    if after is not None:
        sp.set(Q("after"), str(int(after)))
    if line is not None:
        sp.set(Q("line"), str(int(line)))
    if rule is not None:
        sp.set(Q("lineRule"), rule)


def align_abstract_spacing(target: Document):
    """对齐摘要页标题段距与摘要正文行距到书写范例。幂等。"""
    paras = target.paragraphs
    # 找正文区（摘要目录之后）的"摘  要"标题
    toc_seen = False
    abs_title = None
    for p in paras:
        t = p.text.strip().replace(" ", "")
        if t in ("目  录", "目录", "Contents"):
            toc_seen = True
            continue
        if toc_seen and "摘要" in t and "Abstract" not in t:
            if _is_toc_para(p):
                continue  # 跳过目录条目，找真正的摘要标题
            abs_title = p
            break
    if abs_title is None:
        return 0
    # 摘要标题段距
    n = 0
    cur = abs_title
    before_val = cur._p.find(Q("pPr"))
    need = True
    if before_val is not None:
        sp = before_val.find(Q("spacing"))
        if sp is not None and sp.get(Q("before")) == "391" and sp.get(Q("after")) == "312":
            need = False
    if need:
        _set_para_spacing(abs_title, before=391, after=312)
        n += 1
    # 摘要正文（标题之后到 关键词/Keywords 之前）：中文 1.2 倍(288)、英文 1.5 倍(360)
    body_start = False
    for p in paras[paras.index(abs_title) + 1:]:
        t = p.text.strip()
        if not t:
            continue
        if t.startswith("关键词") or t.startswith("Keywords") or t.startswith("Abstract"):
            break
        pPr = p._p.get_or_add_pPr()
        sp = pPr.find(Q("spacing"))
        cur_line = sp.get(Q("line")) if sp is not None else None
        # 摘要正文统一 1.5 倍行距（360），与正文/标题一致；不再区分中英文。
        target_line = 360
        if cur_line != str(target_line):
            _set_para_spacing(p, line=target_line, rule="auto")
            n += 1
    return n


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------
def main():
    ap = argparse.ArgumentParser(description="HIT 硕士论文封面与前置部分版式注入")
    ap.add_argument("docx", help="引擎生成的论文 docx")
    ap.add_argument("--cover-template", default=DEFAULT_COVER, help="官方封面模板 docx")
    ap.add_argument("--front-matter", default=DEFAULT_FRONT, help="front_matter_hit.md 路径")
    ap.add_argument("--cn-title", default=None, help="中文题目（优先于 front_matter）")
    ap.add_argument("--en-title", default=None, help="英文题目（优先于 front_matter）")
    ap.add_argument("--markdown-dir", default=None,
                    help="论文 md 所在目录，用于加载 .env / LLM 翻译配置（与英文目录同一来源）")
    ap.add_argument("--source-md", default=None,
                    help="（预处理后的）论文 md 路径；用于把首行 H1 / 论文题目： 识别为封面题名兜底")
    ap.add_argument("--out", default=None, help="输出路径，默认覆盖原文件")
    ap.add_argument("--no-spacing", action="store_true", help="跳过前置部分行距对齐")
    args = ap.parse_args()

    target = Document(args.docx)
    cn, en = extract_titles(args.front_matter, target, args.cn_title, args.en_title,
                            args.markdown_dir, args.source_md)
    print(f"题目: 中文={cn!r} 英文={en!r}")

    cover_changed = False
    if os.path.exists(args.cover_template):
        print(f"1. 读取封面模板: {args.cover_template}")
        cover = Document(args.cover_template)
        nn = replace_cover_titles(cover, cn, en)
        print(f"   题目替换 {nn} 处")
        elems = build_cover_elems(cover)
        ok = inject_cover(target, elems)
        cover_changed = ok
        print(f"   封面区注入: {'成功' if ok else '跳过'}")
    else:
        print(f"[警告] 未找到封面模板 {args.cover_template}，跳过封面复制（仅做目录/行距）")

    print("2. 目录摘要条目 ...")
    n_toc = ensure_toc_abstract(target)
    print(f"   插入 {n_toc} 条")

    if not args.no_spacing:
        print("3. 前置部分行距对齐 ...")
        n_sp = align_abstract_spacing(target)
        print(f"   调整 {n_sp} 处")

    out = args.out or args.docx
    target.save(out)
    print(f"保存: {out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
