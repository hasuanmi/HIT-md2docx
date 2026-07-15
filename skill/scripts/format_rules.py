#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
HIT 硕士学位论文 · 格式规则引擎（数据驱动，与生成代码解耦）

设计目标（针对"改一个 bug 出另一个"的回归问题）：
- 所有可机检规则集中在本文件，修改/新增规则只动这里，不碰 profile.py 等生成代码。
- 每条规则 = 一个 Rule 对象，含 id / 分类 / 严重度 / 是否可自动修复 / checker / fixer。
- checker 只读，fixer 幂等：修复后再次 audit 该规则必须不再报警（否则测试会失败）。
- audit / fix / report 三个脚本都只 import 本文件，互不耦合。

规则覆盖《硕士学位论文格式要求及审查要点》（哈工大深圳研究生院）高频易错点。
仅做"生成器保证不了、依赖作者 Markdown 内容"的合规项；字体/字号/行距/页眉双线等
已由 hit_master_thesis profile 正确生成，不在此重复校验，避免重复造轮子与回归。
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import Callable, List, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# HIT 自定义样式 id（与 thesis_md2docx/profiles/hit_master_thesis/styles.py 对齐）
# ---------------------------------------------------------------------------
S_BODY = "HitBody"
S_H1 = "HitHeading1"
S_H2 = "HitHeading2"
S_H3 = "HitHeading3"
S_FRONT = "HitFrontHeading"
S_CAPTION = "HitCaption"
S_REFERENCE = "HitReference"
S_FIGURE = "HitFigureImage"
S_TOC_FIELD = "HitTocField"

HEADING_IDS = {S_H1, S_H2, S_H3}
FRONT_HEADING_IDS = {S_FRONT}

# 物理单位词表（用于"数字与单位间加空格"规则）。注意排除了 D/nd 等字母组合。
UNIT_LIST = [
    "cm", "mm", "km", "kg", "g", "mg", "t", "s", "min", "h", "ms", "μs", "ns",
    "MPa", "kPa", "Pa", "GHz", "MHz", "kHz", "Hz", "nm", "μm", "µm", "dB", "W",
    "V", "A", "J", "N", "ml", "mL", "L", "°C", "℃", "rpm", "K", "rad", "bit",
    "B", "KB", "MB", "GB", "bps", "m/s",
]
# 排序：长词优先，避免 "m" 先匹配吃掉 "mm"
UNIT_LIST_SORTED = sorted(UNIT_LIST, key=len, reverse=True)
UNIT_PATTERN = "(" + "|".join(re.escape(u) for u in UNIT_LIST_SORTED) + ")"
# 数字后紧跟可选空格再跟单位，且前后都不是字母/数字（避免 3D / x1 / Figure1）
UNIT_RE = re.compile(
    r"(?<![A-Za-z0-9])((?P<num>\d+)\s*(?P<unit>" + UNIT_PATTERN + r"))(?![A-Za-z0-9])"
)

# 坐标斜杠单位（如 t/s, f/Hz, v/(m·s⁻¹) 等），用于提示改用括号 t(s)
COORD_SLASH_RE = re.compile(r"\b[a-zA-Z][a-zA-Z]*/[a-zA-Z][a-zA-Z]*\b")
# 参考文献上标：间接引用应为上标 [7]；"文献[23]" 直接引用不应上标
DIRECT_QUOTE_RE = re.compile(r"文献\s*\[(\d+)\]")
# 图/表编号（含章号）
FIG_CAP_RE = re.compile(r"^(图|表)\s*(\d+)\s*[-\u2014]\s*(\d+)")
# 图/表编号（缺章号，单数字）
FIG_NO_CHAPTER_RE = re.compile(r"^(图|表)\s*(\d+)\s*$")
# 引用标注
CITE_RE = re.compile(r"\[\d+(?:[\u2010-]|[-])?\d*\]")
# 参考文献条目：以 "[1]" 开头的独立文献行（用于区分“结论里的引用”与“参考文献列表项”）
REFERENCE_ENTRY_RE = re.compile(r"^\[\d+\]")
# 年份
YEAR_RE = re.compile(r"(?:[^\d]|^)(\d{4})(?:[^\d]|$)")
# 英文作者名启发式：含大量拉丁字母且不含明显中文
HAS_CJK = re.compile(r"[\u4e00-\u9fff]")


# ---------------------------------------------------------------------------
# 数据结构
# ---------------------------------------------------------------------------
@dataclass
class Finding:
    rule_id: str
    category: str
    severity: str          # "error" | "warning"
    auto_fixable: bool
    message: str
    location: str = ""     # 段落索引 / 文本片段，便于定位


@dataclass
class Rule:
    id: str
    category: str
    title: str
    detail: str
    severity: str
    auto_fixable: bool
    check: Callable[["ThesisDoc"], List[Finding]]
    fix: Optional[Callable[["ThesisDoc"], int]] = None


# ---------------------------------------------------------------------------
# 文档封装
# ---------------------------------------------------------------------------
class ThesisDoc:
    def __init__(self, path: str):
        self.path = path
        self.doc = Document(path)

    @property
    def paragraphs(self):
        return self.doc.paragraphs

    def paras_with_style(self, *ids):
        return [p for p in self.doc.paragraphs if p.style and p.style.style_id in ids]

    def index_of(self, para) -> int:
        # python-docx 每次访问 .paragraphs 都会生成新的 Paragraph 包装对象，
        # 不能用 list.index(para)（比的是不同 Python 对象）。改为比对底层 lxml 元素。
        target = getattr(para, "_p", None)
        for i, p in enumerate(self.doc.paragraphs):
            if getattr(p, "_p", None) is target:
                return i
        return -1

    def save(self, path: Optional[str] = None):
        self.doc.save(path or self.path)


# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------
def _clear_runs(para):
    for r in list(para.runs):
        r._element.getparent().remove(r._element)


def _set_para_text(para, text: str):
    _clear_runs(para)
    para.add_run(text)


def _is_all_caps(word: str) -> bool:
    """判断是否为全大写缩写（如 PSM-DID、API）：字母部分全为大写即视为专有名词。"""
    letters = [c for c in word if c.isalpha()]
    return bool(letters) and all(c.isupper() for c in letters)


def _set_body_text(para, label: str, body: str):
    """保留标签 runs（如引擎已加粗的“Keywords”/“：”），仅重写标签之后的正文 runs，
    且正文 run 不继承标签的加粗。用于关键词修正，避免 _set_para_text 抹掉加粗标签。"""
    import copy as _copy
    runs = para.runs
    acc = ""
    split_idx = len(runs)
    matched = False
    for i, r in enumerate(runs):
        acc += (r.text or "")
        if acc == label:
            split_idx = i + 1
            matched = True
            break
        if not label.startswith(acc):
            break
    if not matched:
        # 极少数标签未能精确对齐 run 边界，回退为整段重写
        _set_para_text(para, label + body)
        return
    # 移除标签之后的所有 runs
    for r in runs[split_idx:]:
        r._element.getparent().remove(r._element)
    # 参考 rPr：取标签之后首个“非加粗” run（原正文）；否则用引擎默认字号 24（小四）
    ref_rpr = None
    for r in runs[:split_idx]:
        rpr = r._element.find(qn("w:rPr"))
        if rpr is not None and rpr.find(qn("w:b")) is None:
            ref_rpr = rpr
            break
    new_run = para.add_run(body)
    if ref_rpr is not None:
        new_run._element.insert(0, _copy.deepcopy(ref_rpr))
    else:
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "24")
        rPr.append(sz)
        new_run._element.insert(0, rPr)


def _is_english(text: str) -> bool:
    if not text:
        return False
    if HAS_CJK.search(text):
        return False
    letters = sum(1 for c in text if c.isascii() and c.isalpha())
    return letters >= len(text.strip()) * 0.5


# ---------------------------------------------------------------------------
# R01 关键词分隔符与大小写
# ---------------------------------------------------------------------------
def _check_keywords(doc: ThesisDoc) -> List[Finding]:
    out = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("关键词：") or t.startswith("关键词:"):
            body = t.split("：", 1)[-1].split(":", 1)[-1]
            # 中文关键词应用全角分号；若混用半角 ; , 或全角逗号则报警
            if re.search(r"[;,]|\，", body):
                out.append(Finding(
                    "R01", "关键词", "error", True,
                    "中文关键词应使用全角分号“；”分隔，发现半角“;”“,”或全角“，”混用。",
                    f"段落[{doc.index_of(p)}]: {t[:40]}"))
        elif t.startswith("Keywords：") or t.startswith("Keywords:") or t.lower().startswith("keywords"):
            m = re.match(r"^(Keywords\s*[:：]\s*)", t, re.IGNORECASE)
            label = m.group(1) if m else "Keywords: "
            body = t[len(label):]
            # 英文关键词应用逗号分隔且全小写（全大写缩写如 PSM-DID 视为专有名词，不强制小写）
            if "；" in body or ";" in body:
                out.append(Finding(
                    "R01", "关键词", "error", True,
                    "英文关键词应使用半角逗号“,”分隔，发现分号。",
                    f"段落[{doc.index_of(p)}]: {t[:40]}"))
            if any(re.search(r"[A-Z]{2,}", w) and not _is_all_caps(w) for w in re.findall(r"\S+", body)):
                out.append(Finding(
                    "R01", "关键词", "warning", True,
                    "英文关键词应全小写（专有名词/缩写除外）。",
                    f"段落[{doc.index_of(p)}]: {t[:40]}"))
    return out


def _fix_keywords(doc: ThesisDoc) -> int:
    n = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        is_cn = t.startswith("关键词：") or t.startswith("关键词:")
        is_en = t.lower().startswith("keywords")
        if not (is_cn or is_en):
            continue
        # 标签文本（引擎已加粗，需保留）
        if is_cn:
            label = "关键词："
            body = t.split("：", 1)[-1].split(":", 1)[-1]
            fixed_body = re.sub(r"[;,]|\，", "；", body)
            final_text = label + fixed_body
        else:
            m = re.match(r"^(Keywords\s*[:：]\s*)", t, re.IGNORECASE)
            label = m.group(1) if m else "Keywords: "
            body = t[len(label):]
            fixed_body = re.sub(r"[;；\，]", ", ", body)
            # 全大写缩写（PSM-DID 等）保留原样；其余转小写
            fixed_body = ", ".join(
                w.strip() if _is_all_caps(w.strip()) else w.strip().lower()
                for w in fixed_body.split(",")
            )
            final_text = label + fixed_body
        if final_text == t:
            continue
        # 若引擎已加粗标签（首 run 加粗且为“关键词”/“Keywords”），仅就地修正正文，
        # 避免 _set_para_text 把加粗标签一并抹掉。
        runs = p.runs
        already_bold_label = False
        if runs:
            rpr = runs[0]._element.find(qn("w:rPr"))
            already_bold = rpr is not None and rpr.find(qn("w:b")) is not None
            already_bold_label = already_bold and (runs[0].text or "").strip() in ("关键词", "Keywords")
        if already_bold_label:
            _set_body_text(p, label, fixed_body)
        else:
            _set_para_text(p, final_text)
        n += 1
    return n


# ---------------------------------------------------------------------------
# R02 数字与单位间加空格（3cm -> 3 cm）
# ---------------------------------------------------------------------------
def _check_unit_space(doc: ThesisDoc) -> List[Finding]:
    out = []
    for p in doc.paras_with_style(S_BODY, S_CAPTION, S_REFERENCE, S_H1, S_H2, S_H3):
        for m in UNIT_RE.finditer(p.text):
            if m.group(0) != f"{m.group('num')} {m.group('unit')}":
                out.append(Finding(
                    "R02", "排版", "warning", True,
                    f"数字与单位之间应加空格：『{m.group(0)}』建议改为『{m.group('num')} {m.group('unit')}』。",
                    f"段落[{doc.index_of(p)}]"))
                break  # 每段只报一次
    return out


def _fix_unit_space(doc: ThesisDoc) -> int:
    n = 0
    for p in doc.paras_with_style(S_BODY, S_CAPTION, S_REFERENCE, S_H1, S_H2, S_H3):
        changed = False
        for r in p.runs:
            new = UNIT_RE.sub(lambda m: f"{m.group('num')} {m.group('unit')}", r.text)
            if new != r.text:
                r.text = new
                changed = True
        if changed:
            n += 1
    return n


# ---------------------------------------------------------------------------
# R03 图/表编号缺章号（图 1 / 表 1 -> 图 1-1）
# ---------------------------------------------------------------------------
def _check_fig_no_chapter(doc: ThesisDoc) -> List[Finding]:
    out = []
    for p in doc.paras_with_style(S_CAPTION):
        t = p.text.strip()
        m = FIG_NO_CHAPTER_RE.match(t)
        if m:
            out.append(Finding(
                "R03", "图表格公式", "error", False,
                f"图/表编号应包含章号：『{t}』应为『{m.group(1)} {m.group(2)}-1』等形式。",
                f"段落[{doc.index_of(p)}]"))
    return out


# ---------------------------------------------------------------------------
# R04 图/表编号连续性（同章内不应缺号/重号）
# ---------------------------------------------------------------------------
def _check_fig_continuity(doc: ThesisDoc) -> List[Finding]:
    out = []
    caps = {}  # (kind, chapter) -> [numbers]
    for p in doc.paras_with_style(S_CAPTION):
        t = p.text.strip()
        m = FIG_CAP_RE.match(t)
        if m:
            kind, ch, num = m.group(1), int(m.group(2)), int(m.group(3))
            caps.setdefault((kind, ch), []).append((num, doc.index_of(p)))
    for (kind, ch), lst in caps.items():
        nums = sorted(x[0] for x in lst)
        if nums[0] != 1:
            out.append(Finding(
                "R04", "图表格公式", "warning", False,
                f"{kind} {ch} 章编号应从 1 开始，实际从 {nums[0]} 开始。", ""))
        for i in range(1, len(nums)):
            if nums[i] == nums[i - 1]:
                out.append(Finding(
                    "R04", "图表格公式", "warning", False,
                    f"{kind} {ch}-{nums[i]} 编号重复。", ""))
            elif nums[i] > nums[i - 1] + 1:
                out.append(Finding(
                    "R04", "图表格公式", "warning", False,
                    f"{kind} {ch} 章编号缺号：{nums[i-1]} 之后直接到 {nums[i]}。", ""))
    return out


# ---------------------------------------------------------------------------
# R05 直接引用"文献[23]"不应上标
# ---------------------------------------------------------------------------
def _check_direct_quote(doc: ThesisDoc) -> List[Finding]:
    out = []
    for p in doc.paras_with_style(S_BODY, S_H1, S_H2, S_H3):
        for m in DIRECT_QUOTE_RE.finditer(p.text):
            # 找到包含 "[23]" 的 run，检查其上标
            for r in p.runs:
                if f"[{m.group(1)}]" in r.text and getattr(r.font, "superscript", False):
                    out.append(Finding(
                        "R05", "参考文献", "error", True,
                        f"直接引用『文献[{m.group(1)}]』不应使用上标，应与正文排齐。",
                        f"段落[{doc.index_of(p)}]"))
                    break
    return out


def _fix_direct_quote(doc: ThesisDoc) -> int:
    n = 0
    for p in doc.paras_with_style(S_BODY, S_H1, S_H2, S_H3):
        for m in DIRECT_QUOTE_RE.finditer(p.text):
            for r in p.runs:
                if f"[{m.group(1)}]" in r.text and getattr(r.font, "superscript", False):
                    r.font.superscript = False
                    n += 1
                    break
    return n


# ---------------------------------------------------------------------------
# R06 英文摘要页标题应为 Abstract（首字母大写），非 ABSTRACT
# ---------------------------------------------------------------------------
def _check_abstract_title(doc: ThesisDoc) -> List[Finding]:
    out = []
    for p in doc.paras_with_style(*FRONT_HEADING_IDS):
        t = p.text.strip()
        # 英文摘要页标题应恰为“Abstract”（首字母大写）；全大写/全小写均不合规
        if t.lower() == "abstract" and t != "Abstract":
            out.append(Finding(
                "R06", "摘要", "error", True,
                "英文摘要页标题应为首字母大写『Abstract』，而非全大写『ABSTRACT』。",
                f"段落[{doc.index_of(p)}]: {t}"))
    return out


def _fix_abstract_title(doc: ThesisDoc) -> int:
    n = 0
    for p in doc.paras_with_style(*FRONT_HEADING_IDS):
        t = p.text.strip()
        if t.lower() == "abstract" and t != "Abstract":
            _set_para_text(p, "Abstract")
            n += 1
    return n


# ---------------------------------------------------------------------------
# R07 结论中不应出现引用标注 [n]
# ---------------------------------------------------------------------------
def _section_paras(doc: ThesisDoc, heading_text_prefix: str):
    """返回从某一级标题开始到下一个同级标题之前的段落。"""
    paras = doc.paragraphs
    start = None
    for i, p in enumerate(paras):
        if p.style and p.style.style_id == S_H1 and p.text.strip().startswith(heading_text_prefix):
            start = i
            break
    if start is None:
        return []
    end = len(paras)
    for j in range(start + 1, len(paras)):
        if paras[j].style and paras[j].style.style_id == S_H1:
            end = j
            break
    return paras[start:end]


def _check_conclusion_cite(doc: ThesisDoc) -> List[Finding]:
    out = []
    for p in _section_paras(doc, "结论"):
        # 跳过参考文献列表项（"[1] 作者..."），它们可能被误归入结论章范围，
        # 但不是结论正文里的引用标注
        if REFERENCE_ENTRY_RE.match(p.text.strip()):
            continue
        if CITE_RE.search(p.text):
            out.append(Finding(
                "R07", "结论", "error", False,
                "结论部分不应标注引用文献。",
                f"段落[{doc.index_of(p)}]: {p.text.strip()[:40]}"))
    return out


# ---------------------------------------------------------------------------
# R08 标题后直接插图（插图前应有"见图 X-X"提示）
# ---------------------------------------------------------------------------
def _check_heading_before_figure(doc: ThesisDoc) -> List[Finding]:
    out = []
    paras = doc.paragraphs
    i = 0
    while i < len(paras):
        p = paras[i]
        is_heading = p.style and (p.style.style_id in HEADING_IDS or p.style.style_id in FRONT_HEADING_IDS)
        if is_heading:
            j = i + 1
            seen_hint = False
            while j < len(paras):
                q = paras[j]
                if q.style and (q.style.style_id in HEADING_IDS or q.style.style_id in FRONT_HEADING_IDS):
                    break
                if q.style and q.style.style_id == S_FIGURE:
                    if not seen_hint:
                        out.append(Finding(
                            "R08", "图表格公式", "warning", False,
                            "标题后直接插图，插图前应有“见图 X-X / 如图 X-X 所示”的提示文字。",
                            f"标题段落[{doc.index_of(p)}]: {p.text.strip()[:30]}"))
                    break
                if re.search(r"见图|如图|见\s*Figure|as\s+shown\s+in\s+Fig", q.text):
                    seen_hint = True
                j += 1
        i += 1
    return out


# ---------------------------------------------------------------------------
# R09 参考文献数量 / 英文比 / 近五年比
# ---------------------------------------------------------------------------
# 参考文献在规范写法里是顶层一级标题（`# 参考文献`），但部分作者会误写成
# 二级（`## 参考文献`）甚至嵌套在正文章节下。这里兼容 H1/H2 两种层级，
# 并收集其后直到下一个后置章节（致谢/成果/声明/简历/附录）之前的全部文献条目。
_POST_REF_HEADINGS = {
    "参考文献",
    "致谢",
    "个人简历",
    "攻读硕士学位期间取得的科研成果",
    "攻读博士学位期间取得创新性成果",
    "哈尔滨工业大学学位论文原创性声明和使用权限",
    "原创性声明",
}


def _reference_block(doc: ThesisDoc):
    paras = doc.paragraphs
    start = None
    for i, p in enumerate(paras):
        sid = p.style.style_id if p.style else None
        if sid in (S_H1, S_H2) and p.text.strip().startswith("参考文献"):
            start = i
            break
    if start is None:
        return []
    refs = []
    for p in paras[start + 1:]:
        sid = p.style.style_id if p.style else None
        t = p.text.strip()
        # 遇到下一个后置章节标题（非参考文献本身）即停止收集
        if sid in (S_H1, S_H2) and t in _POST_REF_HEADINGS and t != "参考文献":
            break
        if sid == S_REFERENCE or REFERENCE_ENTRY_RE.match(t):
            refs.append(p)
    return refs


def _check_reference_count(doc: ThesisDoc) -> List[Finding]:
    out = []
    refs = _reference_block(doc)
    total = len(refs)
    if total == 0:
        return [Finding("R09", "参考文献", "warning", False,
                        "未检测到参考文献条目（可能目录域未刷新或样式不匹配）。", "")]
    english = sum(1 for p in refs if _is_english(p.text))
    years = []
    for p in refs:
        ms = YEAR_RE.findall(p.text)
        if ms:
            years.append(int(ms[0]))
    import datetime
    cur = datetime.datetime.now().year
    recent5 = sum(1 for y in years if cur - y <= 5)
    recent2 = sum(1 for y in years if cur - y <= 2)
    if total < 40:
        out.append(Finding("R09", "参考文献", "error", False,
                           f"参考文献数量 {total} 篇，少于要求的 40 篇。", ""))
    if english < total / 2:
        out.append(Finding("R09", "参考文献", "warning", False,
                           f"英文文献 {english} 篇（{english/total:.0%}），少于要求的 ≥1/2。", ""))
    if years and recent5 < total / 3:
        out.append(Finding("R09", "参考文献", "warning", False,
                           f"近五年文献 {recent5} 篇（{recent5/total:.0%}），少于要求的 ≥1/3。", ""))
    if years and recent2 == 0:
        out.append(Finding("R09", "参考文献", "warning", False,
                           "未检测到近两年的参考文献。", ""))
    return out


# ---------------------------------------------------------------------------
# R10 坐标单位用斜杠（t/s -> t(s)）
# ---------------------------------------------------------------------------
def _check_coord_slash(doc: ThesisDoc) -> List[Finding]:
    out = []
    for p in doc.paras_with_style(S_BODY, S_CAPTION):
        for m in COORD_SLASH_RE.finditer(p.text):
            tok = m.group(0)
            if "/" in tok and not tok.lower().startswith("http"):
                out.append(Finding(
                    "R10", "图表格公式", "warning", False,
                    f"坐标/物理量单位建议用括号形式，如『{tok}』改为『{tok.split('/')[0]}({tok.split('/')[1]})』。",
                    f"段落[{doc.index_of(p)}]"))
                break
    return out


# ---------------------------------------------------------------------------
# R11 "如下图 X-X 所示：" -> "如图 X-X 所示。"
# ---------------------------------------------------------------------------
def _check_like_figure_colon(doc: ThesisDoc) -> List[Finding]:
    out = []
    pat = re.compile(r"(?<![如图见图])(如下[图表]|如下\s*图|如下\s*表)\s*[\d\u4e00-\u9fff\-]+\s*所示[：:]")
    for p in doc.paras_with_style(S_BODY):
        if pat.search(p.text):
            out.append(Finding(
                "R11", "图表格公式", "warning", False,
                "“如下图 X-X 所示：”应改为“如图 X-X 所示。”（去掉冒号，句末用句号）。",
                f"段落[{doc.index_of(p)}]: {p.text.strip()[:40]}"))
    return out


# ---------------------------------------------------------------------------
# R12 正文英文用 Times New Roman（引擎偶发遗漏的变量/缩写）
# ---------------------------------------------------------------------------
def _body_started(p) -> bool:
    t = p.text.strip()
    return t.startswith("第一章") or t.startswith("Chapter One") or t.startswith("1 ")


def _check_body_tnr(doc: ThesisDoc) -> List[Finding]:
    out = []
    started = False
    for p in doc.paragraphs:
        if not started:
            if _body_started(p):
                started = True
            else:
                continue
        t = p.text.strip()
        if not re.search(r"[A-Za-z]{3,}", t) or t.startswith("Chapter"):
            continue
        for r in p.runs:
            if not re.search(r"[A-Za-z]{3,}", r.text or ""):
                continue
            rf = r._element.find(qn("w:rPr"))
            an = None
            if rf is not None:
                f = rf.find(qn("w:rFonts"))
                an = f.get(qn("w:ascii")) if f is not None else None
            if not (an and "Times New Roman" in an):
                out.append(Finding(
                    "R12", "字体", "warning", True,
                    "正文（变量/缩写/公式中的拉丁字母）应使用 Times New Roman。",
                    f"段落[{doc.index_of(p)}]: {t[:40]}"))
                break
    return out


def _fix_body_tnr(doc: ThesisDoc) -> int:
    n = 0
    started = False
    for p in doc.paragraphs:
        if not started:
            if _body_started(p):
                started = True
            else:
                continue
        if not re.search(r"[A-Za-z0-9]", p.text):
            continue
        for r in p.runs:
            if not re.search(r"[A-Za-z0-9]", r.text or ""):
                continue
            rPr = r._element.get_or_add_rPr()
            rf = rPr.find(qn("w:rFonts"))
            if rf is None:
                rf = OxmlElement("w:rFonts")
                rPr.append(rf)
            an = rf.get(qn("w:ascii"))
            if not (an and "Times New Roman" in an):
                rf.set(qn("w:ascii"), "Times New Roman")
                rf.set(qn("w:hAnsi"), "Times New Roman")
                n += 1
    return n


# ---------------------------------------------------------------------------
# 规则注册表（唯一真相来源）
# ---------------------------------------------------------------------------
RULES: List[Rule] = [
    Rule("R01", "关键词", "关键词分隔符与大小写", "中文用“；”，英文用“,”且全小写。", "error", True,
         _check_keywords, _fix_keywords),
    Rule("R02", "排版", "数字与单位间加空格", "如“3 cm”而非“3cm”。", "warning", True,
         _check_unit_space, _fix_unit_space),
    Rule("R03", "图表格公式", "图/表编号含章号", "应为“图 1-1 / 表 1-1”而非“图 1 / 表 1”。", "error", False,
         _check_fig_no_chapter, None),
    Rule("R04", "图表格公式", "图/表编号连续", "同章内编号不重不漏。", "warning", False,
         _check_fig_continuity, None),
    Rule("R05", "参考文献", "直接引用不上标", "“文献[23]”中 [23] 应与正文排齐，不上标。", "error", True,
         _check_direct_quote, _fix_direct_quote),
    Rule("R06", "摘要", "英文摘要标题大小写", "页标题应为“Abstract”而非“ABSTRACT”。", "error", True,
         _check_abstract_title, _fix_abstract_title),
    Rule("R07", "结论", "结论不标引用", "结论部分不应出现 [n] 引用标注。", "error", False,
         _check_conclusion_cite, None),
    Rule("R08", "图表格公式", "标题后不直插", "插图前应有“见图 X-X”提示。", "warning", False,
         _check_heading_before_figure, None),
    Rule("R09", "参考文献", "参考文献数量与比例", "≥40 篇；英文≥1/2；近5年≥1/3 且含近2年。", "error", False,
         _check_reference_count, None),
    Rule("R10", "图表格公式", "坐标单位用括号", "如 t(s) 而非 t/s。", "warning", False,
         _check_coord_slash, None),
    Rule("R11", "图表格公式", "“如下图”改为“如图”", "“如下图 X-X 所示：”→“如图 X-X 所示。”", "warning", False,
         _check_like_figure_colon, None),
    Rule("R12", "字体", "正文英文用 Times New Roman", "正文变量/缩写/公式中的拉丁字母应为 Times New Roman。", "warning", True,
         _check_body_tnr, _fix_body_tnr),
]

RULES_BY_ID = {r.id: r for r in RULES}


def run_audit(doc: ThesisDoc) -> List[Finding]:
    findings: List[Finding] = []
    for rule in RULES:
        findings.extend(rule.check(doc))
    return findings


def run_fix(doc: ThesisDoc) -> dict:
    """对各规则的 fixer 求幂等修复；返回 {rule_id: 修复计数}。"""
    result = {}
    for rule in RULES:
        if rule.fix:
            result[rule.id] = rule.fix(doc)
    return result


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("usage: format_rules.py <docx>")
        sys.exit(1)
    d = ThesisDoc(sys.argv[1])
    for f in run_audit(d):
        print(f"[{f.severity.upper()}] {f.rule_id} {f.category}: {f.message}")
